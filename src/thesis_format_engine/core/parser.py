import re
from collections import Counter

from docx import Document
from docx.oxml.ns import qn

from thesis_format_engine.models.node import DocumentNode
from thesis_format_engine.models.style import BorderStyle, ParagraphStyleSnapshot, TableStyleSnapshot


class DocxParser:
    def parse(self, path: str) -> list[DocumentNode]:
        document = Document(path)
        nodes: list[DocumentNode] = []

        for index, paragraph in enumerate(document.paragraphs):
            style_name = paragraph.style.name if paragraph.style else None
            text = paragraph.text.strip()
            nodes.append(
                DocumentNode(
                    node_id=f"p-{index}",
                    node_type="paragraph",
                    region=style_name,
                    logical_role=self._infer_logical_role(style_name, text),
                    text=text,
                    paragraph_style=self._extract_paragraph_style(paragraph),
                )
            )

        for index, table in enumerate(document.tables):
            style_name = table.style.name if table.style else "Table"
            nodes.append(
                DocumentNode(
                    node_id=f"t-{index}",
                    node_type="table",
                    region=style_name,
                    logical_role="table",
                    table_style=self._extract_table_style(table),
                )
            )

        return nodes

    def _infer_logical_role(self, style_name: str | None, text: str) -> str:
        normalized = (text or "").strip().lower()
        exact = (text or "").strip()

        if self._is_toc_entry(exact):
            return "toc_entry"
        if self._is_reference_item(exact):
            return "references_item"
        if self._is_numbered_heading(exact):
            return "numbered_heading"

        if style_name == "Heading 1":
            if exact == "摘要" or normalized == "abstract":
                return "abstract_heading"
            if exact == "参考文献":
                return "references_heading"
            return "heading1"
        if style_name == "Heading 2":
            return "heading2"
        if exact == "摘要" or normalized == "abstract":
            return "abstract_heading"
        if exact == "参考文献":
            return "references_heading"
        if exact.startswith("图") and "：" in exact:
            return "figure_caption"
        if exact.startswith("表") and "：" in exact:
            return "table_caption"
        return "body"

    def _is_toc_entry(self, text: str) -> bool:
        return bool(re.match(r"^.+\.{2,}\s*\d+$", text))

    def _is_reference_item(self, text: str) -> bool:
        return bool(re.match(r"^(\[\d+\]|\d+[.、])\s*.+", text))

    def _is_numbered_heading(self, text: str) -> bool:
        return bool(re.match(r"^\d+(?:\.\d+)*\s+.+", text))

    def _extract_paragraph_style(self, paragraph) -> ParagraphStyleSnapshot:
        first_run = paragraph.runs[0] if paragraph.runs else None
        style_font = paragraph.style.font if paragraph.style else None
        fmt = paragraph.paragraph_format

        return ParagraphStyleSnapshot(
            font_family=(first_run.font.name if first_run and first_run.font.name else None)
            or (style_font.name if style_font and style_font.name else None),
            font_size_pt=(first_run.font.size.pt if first_run and first_run.font.size else None)
            or (style_font.size.pt if style_font and style_font.size else None),
            bold=(first_run.font.bold if first_run and first_run.font.bold is not None else None)
            if first_run
            else (style_font.bold if style_font and style_font.bold is not None else None),
            italic=(first_run.font.italic if first_run and first_run.font.italic is not None else None)
            if first_run
            else (style_font.italic if style_font and style_font.italic is not None else None),
            alignment=paragraph.alignment.name if paragraph.alignment is not None else None,
            line_spacing=float(fmt.line_spacing) if fmt.line_spacing is not None else None,
            first_line_indent_pt=fmt.first_line_indent.pt if fmt.first_line_indent is not None else None,
            space_before_pt=fmt.space_before.pt if fmt.space_before is not None else None,
            space_after_pt=fmt.space_after.pt if fmt.space_after is not None else None,
        )

    def _extract_table_style(self, table) -> TableStyleSnapshot:
        tbl_pr = table._tbl.tblPr
        return TableStyleSnapshot(
            alignment=table.alignment.name if table.alignment is not None else None,
            width=self._extract_table_width(tbl_pr),
            top_border=self._extract_border(tbl_pr, "top"),
            bottom_border=self._extract_border(tbl_pr, "bottom"),
            left_border=self._extract_border(tbl_pr, "left"),
            right_border=self._extract_border(tbl_pr, "right"),
            inside_h_border=self._extract_border(tbl_pr, "insideH"),
            inside_v_border=self._extract_border(tbl_pr, "insideV"),
            cell_paragraph_style=self._extract_cell_paragraph_style(table),
        )

    def _extract_cell_paragraph_style(self, table) -> ParagraphStyleSnapshot | None:
        samples = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        samples.append(self._extract_paragraph_style(paragraph).model_dump(exclude_none=True))

        if not samples:
            return None

        merged = {}
        keys = sorted({key for sample in samples for key in sample.keys()})
        for key in keys:
            values = [sample[key] for sample in samples if key in sample]
            if not values:
                continue
            merged[key] = Counter(values).most_common(1)[0][0]

        return ParagraphStyleSnapshot.model_validate(merged) if merged else None

    def _extract_table_width(self, tbl_pr) -> str | None:
        width = tbl_pr.find(qn("w:tblW")) if tbl_pr is not None else None
        if width is None:
            return None
        value = width.get(qn("w:w"))
        width_type = width.get(qn("w:type"))
        if value is None:
            return None
        return f"{value}:{width_type or 'unknown'}"

    def _extract_border(self, tbl_pr, side: str) -> BorderStyle | None:
        if tbl_pr is None:
            return None
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            return None
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            return None
        return BorderStyle(
            style=node.get(qn("w:val")),
            size=int(node.get(qn("w:sz"))) if node.get(qn("w:sz")) else None,
            color=node.get(qn("w:color")),
        )
