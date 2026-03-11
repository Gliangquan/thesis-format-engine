import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from thesis_format_engine.models.rule import RuleItem, RuleSet


class PatchEngine:
    def apply(self, docx_path: str, rules: RuleSet, output_path: str) -> dict:
        document = Document(docx_path)
        changes = 0

        for paragraph in document.paragraphs:
            region = paragraph.style.name if paragraph.style else None
            logical_role = self._infer_logical_role(region, paragraph.text.strip())
            rule = self._match_rule(region, logical_role, rules.regions)
            if not rule or not rule.paragraph_style:
                continue
            changes += self._patch_paragraph(paragraph, rule.paragraph_style.model_dump(exclude_none=True))

        for table in document.tables:
            region = table.style.name if table.style else "Table"
            rule = self._match_rule(region, "table", rules.regions)
            if not rule or not rule.table_style:
                continue
            changes += self._patch_table(table, rule.table_style.model_dump(exclude_none=True))

        document.save(output_path)
        return {"output": output_path, "changes": changes}

    def _match_rule(self, region: str | None, logical_role: str | None, rules: list[RuleItem]) -> RuleItem | None:
        for rule in rules:
            if rule.logical_role and rule.logical_role == logical_role:
                return rule
        for rule in rules:
            if rule.region and rule.region == region:
                return rule
        return None

    def _infer_logical_role(self, style_name: str | None, text: str) -> str:
        normalized = (text or "").strip().lower()
        exact = (text or "").strip()

        if self._is_toc_entry(exact):
            return "toc_entry"
        if self._is_reference_item(exact):
            return "references_item"
        if self._is_numbered_heading(exact):
            return "numbered_heading"

        if style_name == "Heading 1":
            if exact == "摘要" or normalized == "abstract":
                return "abstract_heading"
            if exact == "参考文献":
                return "references_heading"
            return "heading1"
        if style_name == "Heading 2":
            return "heading2"
        if exact == "摘要" or normalized == "abstract":
            return "abstract_heading"
        if exact == "参考文献":
            return "references_heading"
        if exact.startswith("图") and "：" in exact:
            return "figure_caption"
        if exact.startswith("表") and "：" in exact:
            return "table_caption"
        return "body"

    def _is_toc_entry(self, text: str) -> bool:
        return bool(re.match(r"^.+\.{2,}\s*\d+$", text))

    def _is_reference_item(self, text: str) -> bool:
        return bool(re.match(r"^(\[\d+\]|\d+[.、])\s*.+", text))

    def _is_numbered_heading(self, text: str) -> bool:
        return bool(re.match(r"^\d+(?:\.\d+)*\s+.+", text))

    def _patch_paragraph(self, paragraph, expected: dict) -> int:
        changes = 0
        for run in paragraph.runs:
            if "font_family" in expected and run.font.name != expected["font_family"]:
                run.font.name = expected["font_family"]
                changes += 1
            if "font_size_pt" in expected:
                target_size = float(expected["font_size_pt"])
                current = run.font.size.pt if run.font.size else None
                if current != target_size:
                    run.font.size = self._pt(target_size)
                    changes += 1
            if "bold" in expected and run.font.bold != expected["bold"]:
                run.font.bold = expected["bold"]
                changes += 1

        if "alignment" in expected:
            current_alignment = paragraph.alignment.name if paragraph.alignment is not None else None
            if current_alignment != expected["alignment"]:
                paragraph.alignment = self._alignment(expected["alignment"])
                changes += 1

        fmt = paragraph.paragraph_format
        if "line_spacing" in expected and fmt.line_spacing != expected["line_spacing"]:
            fmt.line_spacing = float(expected["line_spacing"])
            changes += 1
        if "first_line_indent_pt" in expected:
            current_indent = fmt.first_line_indent.pt if fmt.first_line_indent is not None else None
            if current_indent != expected["first_line_indent_pt"]:
                fmt.first_line_indent = self._pt(float(expected["first_line_indent_pt"]))
                changes += 1
        if "space_before_pt" in expected:
            current_space_before = fmt.space_before.pt if fmt.space_before is not None else None
            if current_space_before != expected["space_before_pt"]:
                fmt.space_before = self._pt(float(expected["space_before_pt"]))
                changes += 1
        if "space_after_pt" in expected:
            current_space_after = fmt.space_after.pt if fmt.space_after is not None else None
            if current_space_after != expected["space_after_pt"]:
                fmt.space_after = self._pt(float(expected["space_after_pt"]))
                changes += 1

        return changes

    def _patch_table(self, table, expected: dict) -> int:
        changes = 0
        border_mapping = {
            "top_border": "top",
            "bottom_border": "bottom",
            "left_border": "left",
            "right_border": "right",
            "inside_h_border": "insideH",
            "inside_v_border": "insideV",
        }

        for field, edge in border_mapping.items():
            if field not in expected:
                continue
            self._set_table_border(table, edge, expected[field])
            changes += 1

        if "cell_paragraph_style" in expected:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        changes += self._patch_paragraph(paragraph, expected["cell_paragraph_style"])

        return changes

    def _set_table_border(self, table, edge: str, spec: dict) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)

        border = tbl_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tbl_borders.append(border)

        if "style" in spec:
            border.set(qn("w:val"), str(spec["style"]))
        if "size" in spec:
            border.set(qn("w:sz"), str(spec["size"]))
        if "color" in spec:
            border.set(qn("w:color"), str(spec["color"]))

    def _alignment(self, value: str):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        return getattr(WD_ALIGN_PARAGRAPH, value)

    def _pt(self, value: float):
        from docx.shared import Pt

        return Pt(value)
